"""
Media Processor — Multimodal input pipeline for Tiểu Bảo Bảo AI Agent.

Handles three input modalities received via Telegram:
  1. Voice / Audio  → Groq Whisper STT (turbo → v3 fallback)
  2. Image / Photo  → Groq Vision (qwen3.8-27b → qwen3.6-27b)
                      → OpenRouter Vision (gemma-4-31b:free → gemma-4-26b:free → openrouter/free)
  3. Document       → Local offline extraction (PyMuPDF / python-docx / openpyxl / builtin)
"""

import base64
import csv
import io
import json
import logging
import os
import shutil
import subprocess
import tempfile
from pathlib import Path
from typing import List, Optional, Tuple

import httpx
try:
    from PIL import Image
except ImportError:
    Image = None  # type: ignore

from app.config import settings

logger = logging.getLogger(__name__)

# Max document/archive text fed into LLM context.
# Raised to 12000 to handle archives with up to 3 images (~2500 chars/image description each)
_MAX_DOC_CHARS = 12_000
# Max image dimension before resizing (Groq 4MB base64 limit)
_MAX_IMG_DIMENSION = 1024
_IMG_QUALITY = 85

# OpenRouter fallback chain for vision (verified free tier, Aug 2026)
_OR_VISION_MODELS = [
    settings.OPENROUTER_VISION_MODEL,           # google/gemma-4-31b-it:free
    settings.OPENROUTER_VISION_MODEL_FALLBACK,  # google/gemma-4-26b-a4b-it:free
    "openrouter/free",                          # auto-router last resort
]

_OR_VISION_HEADERS_BASE = {
    "HTTP-Referer": "https://dashboard.kirito.server",
    "X-Title": "Tiểu Bảo Bảo AI Agent",
    "Content-Type": "application/json",
}

_VISION_SYSTEM_PROMPT = (
    "Bạn là trợ lý AI thông minh tên Tiểu Bảo Bảo. "
    "Hãy phân tích hình ảnh được cung cấp và mô tả chi tiết, chính xác nội dung. "
    "Trả lời hoàn toàn bằng tiếng Việt, ngắn gọn và súc tích."
)

# Technical Extraction Prompt for images inside archives/documents:
# Must be concise, purely factual (metrics, charts, numbers, text), with ZERO greetings or conclusions.
_VISION_ARCHIVE_PROMPT = (
    "Bạn là hệ thống trích xuất dữ liệu hình ảnh kỹ thuật cao. "
    "Nhiệm vụ: Trích xuất chính xác, súc tích toàn bộ thông tin quan trọng trong ảnh: "
    "tiêu đề, tên mô hình/dự án, biểu đồ (trục X, trục Y, xu hướng đường cong), các chỉ số metrics (loss, accuracy, RMSE, F1,...), "
    "bảng số liệu (cột, giá trị các hàng nổi bật) và trạng thái. "
    "Quy tắc tối quan trọng: "
    "1. TUYỆT ĐỐI KHÔNG viết lời chào mở đầu (như 'Chào bạn', 'Tôi là Tiểu Bảo Bảo'). "
    "2. TUYỆT ĐỐI KHÔNG viết kết luận chung chung hay lời kết bài. "
    "3. Trình bày bằng các gạch đầu dòng súc tích, ngắn gọn, đi thẳng vào các số liệu và sự thật kỹ thuật. "
    "4. Viết 100% bằng tiếng Việt."
)


class ArchivePasswordRequiredError(ValueError):
    """Raised when an archive is encrypted and requires a password to extract."""
    pass


class ArchiveInvalidPasswordError(ValueError):
    """Raised when the provided password fails to decrypt the archive."""
    pass


class ArchiveCorruptedError(ValueError):
    """Raised when an archive file is malformed or corrupted."""
    pass


def extract_password_from_text(text: str) -> Optional[str]:
    """
    Extracts archive password from user caption or message using smart patterns.
    Examples:
      - 'pass: 123456' -> '123456'
      - 'pass 123456' -> '123456'
      - 'mật khẩu là: Abc@123' -> 'Abc@123'
      - 'mật khẩu là 123' -> '123'
      - 'mật khẩu 123' -> '123'
      - 'mk: test123' -> 'test123'
      - 'mk 123' -> '123'
      - 'password = mypass' -> 'mypass'
      - 'pass là "secret key"' -> 'secret key'
    """
    if not text:
        return None
    import re
    # 1. Match patterns with quotes: pass: "my secret password", pass là '123 456'
    quoted_match = re.search(
        r'(?:pass(?:word)?|mật\s*khẩu|mk)(?:\s*(?:là|is))?\s*[:=]?\s*["\']([^"\']+)["\']',
        text,
        re.IGNORECASE
    )
    if quoted_match:
        return quoted_match.group(1).strip()

    # 2. Match patterns without quotes: pass: 123456, pass 123, mk abc, mk 123, mật khẩu là: 123, mật khẩu 123
    unquoted_match = re.search(
        r'(?:pass(?:word)?|mật\s*khẩu|mk)(?:\s*(?:là|is))?\s*[:=]?\s*([^\s\n,;]+)',
        text,
        re.IGNORECASE
    )
    if unquoted_match:
        val = unquoted_match.group(1).strip()
        if val and val not in (":", "=", "là", "is"):
            return val

    return None


class MediaProcessor:
    """
    Processes multimodal inputs from Telegram: voice, photo, document.
    Designed to be injected into TelegramBot and reuse its httpx.AsyncClient.
    """

    def __init__(self, http_client: Optional[httpx.AsyncClient] = None) -> None:
        # Reuse caller's client if provided to avoid extra connection pools
        self._http = http_client or httpx.AsyncClient(timeout=60.0)
        self._owns_client = http_client is None

    async def aclose(self) -> None:
        if self._owns_client:
            await self._http.aclose()

    # ─── Step 1: Download file from Telegram ─────────────────────────────────

    async def download_telegram_file(self, file_id: str) -> bytes:
        """
        Two-step Telegram file download:
        1. getFile → resolve file_path
        2. Stream download from file CDN URL
        Max file size allowed by Telegram Bot API: 20 MB.
        """
        token = settings.TELEGRAM_BOT_TOKEN
        get_file_url = f"https://api.telegram.org/bot{token}/getFile"
        resp = await self._http.get(get_file_url, params={"file_id": file_id})
        resp.raise_for_status()
        data = resp.json()
        if not data.get("ok"):
            raise ValueError(f"Telegram getFile failed: {data.get('description')}")
        file_path = data["result"]["file_path"]

        cdn_url = f"https://api.telegram.org/file/bot{token}/{file_path}"
        dl_resp = await self._http.get(cdn_url, timeout=60.0)
        dl_resp.raise_for_status()
        return dl_resp.content

    # ─── Modality 1: Voice → Groq Whisper STT ────────────────────────────────

    # Known Whisper hallucination patterns (trained on YouTube/podcast data).
    # Whisper "fills in" these phrases when audio is short, quiet, or ambiguous.
    # We reject transcriptions that are ONLY these patterns to avoid false positives.
    _HALLUCINATION_BLACKLIST = [
        # Vietnamese YouTube/livestream patterns
        "cảm ơn các bạn đã theo dõi",
        "đăng ký kênh",
        "subscribe",
        "hẹn gặp lại",
        "like và subscribe",
        "nhấn chuông thông báo",
        "cảm ơn các bạn",
        "xin chào các bạn",
        "chúc các bạn",
        "ghiền mì gõ",
        "cảm ơn bạn đã xem",
        # English YouTube patterns
        "thank you for watching",
        "please subscribe",
        "don't forget to subscribe",
        "see you in the next video",
        "like and subscribe",
        "hit the bell",
        # Common noise/silence patterns
        ".",
        " .",
        "...",
        "thank you",
        "thanks for watching",
    ]

    @staticmethod
    def _is_hallucination(text: str, duration_seconds: int) -> bool:
        """
        Detects Whisper hallucinations by checking:
        1. Known hallucination phrases (YouTube/livestream patterns)
        2. Suspiciously long transcription for very short audio (ratio check)
        """
        if not text:
            return False
        t_lower = text.lower().strip()

        # Check against known hallucination patterns
        for pattern in MediaProcessor._HALLUCINATION_BLACKLIST:
            if pattern in t_lower:
                # Only reject if the hallucination pattern DOMINATES the transcription
                # (i.e., the pattern makes up >60% of total text length)
                if len(pattern) / max(len(t_lower), 1) > 0.6:
                    return True

        # Suspiciously verbose for very short audio (< 4s producing > 60 chars is suspect)
        if duration_seconds <= 3 and len(text) > 60:
            return True

        return False

    async def transcribe_voice(
        self,
        audio_bytes: bytes,
        filename: str = "voice.oga",
        language: str = "vi",
        duration: int = 0,
    ) -> str:
        """
        Sends audio to Groq Whisper STT.
        Primary: whisper-large-v3-turbo (216x real-time, 20 RPM, 28,800 audio-sec/day).
        Fallback: whisper-large-v3.

        Anti-hallucination measures:
        - temperature=0: greedy decoding — no random sampling → far fewer hallucinations
        - prompt anchor: primes model toward conversational Vietnamese (away from YouTube patterns)
        - blacklist filter: rejects known hallucination phrases post-transcription

        NOTE: Groq Whisper rejects .oga extension (Telegram native format).
        We rename to .ogg which Groq accepts — same OGG container, Opus codec.
        """
        groq_keys = settings.groq_keys
        if not groq_keys:
            return ""

        # Groq whitelist: flac, mp3, mp4, mpeg, mpga, m4a, ogg, opus, wav, webm
        # Telegram voice messages use .oga (OGG Opus) which is NOT in Groq's whitelist.
        ext = Path(filename).suffix.lower()
        if ext == ".oga":
            filename = Path(filename).with_suffix(".ogg").name
        _GROQ_AUDIO_EXTS = {".flac", ".mp3", ".mp4", ".mpeg", ".mpga", ".m4a", ".ogg", ".opus", ".wav", ".webm"}
        if Path(filename).suffix.lower() not in _GROQ_AUDIO_EXTS:
            filename = "voice.ogg"

        # Whisper prompt acts as a "fictitious transcript prefix" — the model predicts
        # the next token after the prompt. Only the final 224 tokens of prompt matter.
        # Test result (Aug 2026 on file 6091204043177206473.ogg "chào bạn"):
        #   ✅ "bạn, chào bạn, xin chào bạn"  → correct: 'chào bạn'
        #   ❌ "Cuộc trò chuyện bằng tiếng Việt: ..."  → wrong: 'Chào bác'
        # Lesson: SHORT, FOCUSED vocabulary prompts beat instruction-style prompts.
        # "bạn" biases away from the acoustically similar "bác" (tonal confusion).
        _STT_PROMPT = "bạn, chào bạn, xin chào bạn, bạn ơi, cảm ơn bạn, được rồi, ổn rồi."

        models = [settings.GROQ_WHISPER_MODEL, settings.GROQ_WHISPER_MODEL_FALLBACK]

        for model in models:
            for key in groq_keys[:3]:
                try:
                    resp = await self._http.post(
                        "https://api.groq.com/openai/v1/audio/transcriptions",
                        headers={"Authorization": f"Bearer {key}"},
                        data={
                            "model": model,
                            "language": language,
                            "response_format": "json",
                            "temperature": "0",    # greedy decoding → minimal hallucination
                            "prompt": _STT_PROMPT, # vocabulary anchor → tonal accuracy
                        },
                        files={"file": (filename, audio_bytes, "audio/ogg")},
                        timeout=90.0,
                    )
                    if resp.status_code == 200:
                        text = resp.json().get("text", "").strip()
                        # Post-processing: reject known hallucination patterns
                        if self._is_hallucination(text, duration):
                            logger.warning(
                                "[MediaProcessor] STT hallucination detected (model=%s): '%s' → discarding",
                                model, text[:80]
                            )
                            return ""
                        logger.info("[MediaProcessor] STT ✅ model=%s len=%d text='%s'", model, len(text), text[:60])
                        return text
                    if resp.status_code == 429:
                        logger.warning("[MediaProcessor] STT 429 key=%s... model=%s", key[:8], model)
                        continue
                    logger.warning("[MediaProcessor] STT HTTP %d: %s", resp.status_code, resp.text[:200])
                except Exception as exc:
                    logger.error("[MediaProcessor] STT error: %s", exc)

        logger.error("[MediaProcessor] STT failed — all Groq keys exhausted")
        return ""

    # ─── Modality 2: Image → Vision (Groq → OpenRouter fallback) ─────────────

    @staticmethod
    def _encode_image(image_bytes: bytes) -> str:
        """Resize to max 1024px JPEG and encode as base64 to keep payload < 4MB."""
        if Image is not None:
            with Image.open(io.BytesIO(image_bytes)) as img:
                img = img.convert("RGB")
                img.thumbnail((_MAX_IMG_DIMENSION, _MAX_IMG_DIMENSION), Image.LANCZOS)
                buf = io.BytesIO()
                img.save(buf, format="JPEG", quality=_IMG_QUALITY)
                return base64.b64encode(buf.getvalue()).decode("utf-8")
        return base64.b64encode(image_bytes).decode("utf-8")

    @staticmethod
    def _build_vision_messages(b64: str, caption: str, is_archive: bool = False) -> list:
        sys_prompt = _VISION_ARCHIVE_PROMPT if is_archive else _VISION_SYSTEM_PROMPT
        if caption:
            user_text = caption
        elif is_archive:
            user_text = "Trích xuất và tóm tắt toàn bộ dữ liệu kỹ thuật, bảng số liệu, biểu đồ, văn bản trong hình ảnh này một cách súc tích, ngắn gọn."
        else:
            user_text = "Hãy phân tích và mô tả chi tiết nội dung hình ảnh này."

        return [
            {"role": "system", "content": sys_prompt},
            {
                "role": "user",
                "content": [
                    {"type": "text", "text": user_text},
                    {"type": "image_url", "image_url": {"url": f"data:image/jpeg;base64,{b64}"}},
                ],
            },
        ]

    async def _analyze_image_groq(self, b64: str, caption: str, is_archive: bool = False) -> Optional[str]:
        """
        Try Groq Vision: qwen3.8-27b → qwen3.6-27b.
        Note: tool_calls not passed — Groq vision runs in analysis-only mode,
        result is then forwarded to ai_agent.chat() for full ReAct processing.
        """
        groq_keys = settings.groq_keys
        if not groq_keys:
            return None

        models = [settings.GROQ_VISION_MODEL, settings.GROQ_VISION_MODEL_FALLBACK]
        messages = self._build_vision_messages(b64, caption, is_archive=is_archive)
        # Archive images must be compact to avoid 413 when the full payload
        # (2-3 image descriptions + system prompt + history) is sent to the main LLM
        max_tokens = 500 if is_archive else 1024

        for model in models:
            for key in groq_keys[:3]:
                try:
                    resp = await self._http.post(
                        "https://api.groq.com/openai/v1/chat/completions",
                        headers={"Authorization": f"Bearer {key}", "Content-Type": "application/json"},
                        json={"model": model, "messages": messages, "temperature": 0.2, "max_tokens": max_tokens},
                        timeout=60.0,
                    )
                    if resp.status_code == 200:
                        content = resp.json()["choices"][0]["message"]["content"]
                        logger.info("[MediaProcessor] Vision-Groq ✅ model=%s", model)
                        return content.strip()
                    if resp.status_code == 429:
                        logger.warning("[MediaProcessor] Vision-Groq 429 key=%s... model=%s", key[:8], model)
                        continue
                    logger.warning("[MediaProcessor] Vision-Groq HTTP %d: %s", resp.status_code, resp.text[:150])
                except Exception as exc:
                    logger.error("[MediaProcessor] Vision-Groq error: %s", exc)

        return None

    async def _analyze_image_openrouter(self, b64: str, caption: str, is_archive: bool = False) -> Optional[str]:
        """
        OpenRouter Vision fallback chain:
        gemma-4-31b:free → gemma-4-26b:free → openrouter/free (auto-router).
        These models support vision + tool_calls simultaneously.
        """
        or_keys = settings.openrouter_keys
        if not or_keys:
            return None

        messages = self._build_vision_messages(b64, caption, is_archive=is_archive)
        max_tokens = 500 if is_archive else 1024

        for model in _OR_VISION_MODELS:
            for key in or_keys[:2]:
                try:
                    resp = await self._http.post(
                        settings.OPENROUTER_API_URL,
                        headers={**_OR_VISION_HEADERS_BASE, "Authorization": f"Bearer {key}"},
                        json={"model": model, "messages": messages, "temperature": 0.2, "max_tokens": max_tokens},
                        timeout=60.0,
                    )
                    if resp.status_code == 200:
                        content = resp.json()["choices"][0]["message"]["content"]
                        logger.info("[MediaProcessor] Vision-OR ✅ model=%s", model)
                        return content.strip()
                    if resp.status_code == 429:
                        logger.warning("[MediaProcessor] Vision-OR 429 key=%s... model=%s", key[:8], model)
                        continue
                    logger.warning("[MediaProcessor] Vision-OR HTTP %d: %s", resp.status_code, resp.text[:150])
                except Exception as exc:
                    logger.error("[MediaProcessor] Vision-OR error: %s", exc)

        return None

    async def analyze_image(self, image_bytes: bytes, caption: str = "", is_archive: bool = False) -> str:
        """
        Orchestrator: Groq Vision → OpenRouter Vision → graceful text fallback.
        Returns a string description to be prepended and passed to ai_agent.chat().
        """
        try:
            b64 = self._encode_image(image_bytes)
        except Exception as exc:
            logger.error("[MediaProcessor] Image encode error: %s", exc)
            return "(Không thể đọc định dạng ảnh này)"

        result = await self._analyze_image_groq(b64, caption, is_archive=is_archive)
        if not result:
            logger.info("[MediaProcessor] Groq vision failed, trying OpenRouter...")
            result = await self._analyze_image_openrouter(b64, caption, is_archive=is_archive)

        if result:
            return result

        logger.error("[MediaProcessor] All vision providers failed")
        return (
            "(Em nhận được ảnh nhưng chưa thể phân tích lúc này do rate limit. "
            "Anh Mạnh vui lòng mô tả nội dung ảnh để em hỗ trợ.)"
        )

    # ─── Modality 3: Universal Document & Archive Extraction (Offline & Multimodal) ───

    _IMAGE_EXTENSIONS = frozenset({
        ".jpg", ".jpeg", ".png", ".gif", ".bmp", ".webp", ".tiff", ".tif",
    })

    _TEXT_EXTENSIONS = frozenset({
        ".txt", ".md", ".py", ".js", ".ts", ".sh", ".bash",
        ".yaml", ".yml", ".toml", ".ini", ".env", ".log",
        ".html", ".xml", ".css", ".sql", ".rs", ".go",
        ".java", ".c", ".cpp", ".h", ".hpp", ".cs", ".php",
        ".rb", ".swift", ".kt", ".kts", ".scala", ".r",
        ".conf", ".cfg", ".properties", ".bat", ".ps1",
    })

    @staticmethod
    def _extract_pdf(data: bytes) -> str:
        try:
            import pymupdf as fitz
        except ImportError:
            import fitz
        pages: List[str] = []
        with fitz.open(stream=data, filetype="pdf") as doc:
            for page_num in range(min(len(doc), 50)):
                text = doc[page_num].get_text("text").strip()
                if text:
                    pages.append(f"--- [Trang {page_num + 1}] ---\n{text}")
        return "\n\n".join(pages)

    @staticmethod
    def _extract_docx(data: bytes) -> str:
        import docx
        doc = docx.Document(io.BytesIO(data))
        parts: List[str] = []
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)
        for idx, table in enumerate(doc.tables, 1):
            parts.append(f"\n[Bảng {idx}]:")
            for row in table.rows:
                parts.append(" | ".join(cell.text.strip().replace("\n", " ") for cell in row.cells))
        return "\n".join(parts)

    @staticmethod
    def _extract_xlsx(data: bytes) -> str:
        import openpyxl
        wb = openpyxl.load_workbook(io.BytesIO(data), data_only=True, read_only=True)
        sheets: List[str] = []
        for name in wb.sheetnames:
            ws = wb[name]
            rows: List[str] = [f"=== Sheet: {name} ==="]
            for row_idx, row in enumerate(ws.iter_rows(values_only=True), 1):
                if row_idx > 200:
                    rows.append("... (cắt bớt sau 200 dòng)")
                    break
                if any(v is not None for v in row):
                    rows.append(" | ".join(str(v) if v is not None else "" for v in row))
            if len(rows) > 1:
                sheets.append("\n".join(rows))
        wb.close()
        return "\n\n".join(sheets)

    @staticmethod
    def _extract_csv(data: bytes) -> str:
        text = data.decode("utf-8", errors="replace")
        reader = csv.reader(io.StringIO(text))
        rows: List[str] = []
        for idx, row in enumerate(reader, 1):
            if idx > 200:
                rows.append("... (cắt bớt sau 200 dòng)")
                break
            rows.append(" | ".join(row))
        return "\n".join(rows)

    @staticmethod
    def _extract_json(data: bytes) -> str:
        parsed = json.loads(data.decode("utf-8", errors="replace"))
        return json.dumps(parsed, indent=2, ensure_ascii=False)

    @classmethod
    def _extract_child_file_content(cls, filename: str, data: bytes) -> tuple[str, str, Optional[str]]:
        """
        Universal extractor for individual files inside or outside archives.
        Returns: (icon, type_label, text_content_or_None)
        """
        ext = Path(filename).suffix.lower()

        try:
            if ext == ".pdf":
                return "📄", "Tài liệu PDF", cls._extract_pdf(data)
            if ext in (".docx", ".doc"):
                return "📄", "Tài liệu Word", cls._extract_docx(data)
            if ext in (".xlsx", ".xlsm", ".xls"):
                return "📊", "Bảng tính Excel", cls._extract_xlsx(data)
            if ext == ".csv":
                return "📊", "Dữ liệu CSV", cls._extract_csv(data)
            if ext == ".json":
                return "📋", "Dữ liệu JSON", cls._extract_json(data)
            if ext in cls._TEXT_EXTENSIONS:
                return "📝", "Mã nguồn / Văn bản", data.decode("utf-8", errors="replace").strip()
            if ext in cls._IMAGE_EXTENSIONS:
                return "🖼️", "Hình ảnh", None  # Will be processed via Vision API
            return "⚙️", "Tệp nhị phân", None
        except Exception as exc:
            logger.warning("[MediaProcessor] Error extracting child file %s: %s", filename, exc)
            return "⚠️", "Lỗi đọc tệp", f"(Không thể trích xuất nội dung: {exc})"

    @staticmethod
    def _detect_archive_type(ext: str, mime: str, data: bytes) -> str:
        """
        Robust 3-tier archive format detector:
        1. File extension (highest confidence)
        2. Magic bytes (ground truth)
        3. MIME type (fallback)
        """
        ext = ext.lower()
        if ext in (".zip", ".jar", ".war"):
            return "zip"
        if ext in (".rar",):
            return "rar"
        if ext in (".tar", ".gz", ".tgz", ".bz2", ".tbz2", ".xz", ".txz") or ext.endswith((".tar.gz", ".tar.bz2", ".tar.xz")):
            return "tar"
        if ext in (".7z",):
            return "7z"

        # Magic bytes inspection
        if data[:4] == b"PK\x03\x04":
            return "zip"
        if data[:7] in (b"Rar!\x1a\x07\x00", b"Rar!\x1a\x07\x01"):
            return "rar"
        if data[:6] == b"7z\xbc\xaf\x27\x1c":
            return "7z"
        if data[:2] == b"\x1f\x8b" or data[:3] == b"BZh" or data[:6] == b"\xfd7zXZ\x00":
            return "tar"

        # MIME fallback
        if "zip" in mime:
            return "zip"
        if "rar" in mime or "x-rar" in mime:
            return "rar"
        if "tar" in mime or "gzip" in mime or "bzip2" in mime:
            return "tar"
        if "7z" in mime:
            return "7z"

        return "unknown"

    @classmethod
    def _unpack_via_7z(
        cls,
        file_bytes: bytes,
        filename: str,
        password: Optional[str] = None,
    ) -> list[tuple[str, int, bytes]]:
        """
        High-performance extraction using 7-Zip CLI (p7zip-full).
        Supports ZIP (ZipCrypto, WinZip AES-256), RAR4, RAR5, 7Z, TAR, GZ, BZ2, XZ.
        """
        seven_zip = shutil.which("7z") or shutil.which("7za")
        if not seven_zip:
            raise FileNotFoundError("7z binary not found on host system")

        _MAX_TOTAL_BYTES = 50 * 1024 * 1024
        _MAX_FILES = 200
        _MAX_SINGLE_FILE = 10 * 1024 * 1024

        with tempfile.TemporaryDirectory(prefix="tbb_archive_") as tmp_dir:
            tmp_path = Path(tmp_dir)
            ext = Path(filename).suffix or ".zip"
            archive_file = tmp_path / f"archive{ext}"
            archive_file.write_bytes(file_bytes)
            extract_dir = tmp_path / "extracted"
            extract_dir.mkdir(parents=True, exist_ok=True)

            # Step 1: Probe archive headers using `7z l -slt -p-` to detect encryption
            probe_cmd = [seven_zip, "l", "-slt", "-p-", str(archive_file)]
            probe_proc = subprocess.run(
                probe_cmd,
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
                text=True,
                timeout=15,
            )
            probe_out = (probe_proc.stdout or "") + "\n" + (probe_proc.stderr or "")

            is_encrypted = (
                "Encrypted = +" in probe_out
                or "7zAES" in probe_out
                or "AES" in probe_out
                or "ZipCrypto" in probe_out
                or "Errors: 1" in probe_out
                or probe_proc.returncode != 0
            )

            if is_encrypted and not password:
                raise ArchivePasswordRequiredError(f"Tệp nén `{filename}` được đặt mật khẩu bảo vệ.")

            # Step 2: Execute extraction with password or -p-
            cmd = [seven_zip, "x", "-y", f"-o{extract_dir}"]
            if password is not None:
                cmd.append(f"-p{password}")
            else:
                cmd.append("-p-")  # Do not prompt interactively

            cmd.append(str(archive_file))

            proc = subprocess.run(
                cmd,
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
                text=True,
                timeout=30,
            )

            out = (proc.stdout or "") + "\n" + (proc.stderr or "")
            out_lower = out.lower()

            if proc.returncode != 0:
                if is_encrypted or password:
                    if password:
                        raise ArchiveInvalidPasswordError(f"Mật khẩu '{password}' không chính xác cho tệp `{filename}`.")
                    raise ArchivePasswordRequiredError(f"Tệp nén `{filename}` được đặt mật khẩu bảo vệ.")

                raise ArchiveCorruptedError(f"Lỗi khi đọc tệp nén `{filename}`: {out.strip()[:180]}")

            members_data: list[tuple[str, int, bytes]] = []
            total_bytes = 0

            for root, _, files in os.walk(extract_dir):
                for f in files:
                    full_path = Path(root) / f
                    rel_name = str(full_path.relative_to(extract_dir)).replace("\\", "/")

                    # Strict Zip Slip guard
                    resolved = (extract_dir / rel_name).resolve()
                    if not str(resolved).startswith(str(extract_dir.resolve())):
                        continue

                    f_size = full_path.stat().st_size
                    if f_size > _MAX_SINGLE_FILE:
                        continue

                    total_bytes += f_size
                    if total_bytes > _MAX_TOTAL_BYTES:
                        raise ValueError(f"Dung lượng giải nén quá lớn ({total_bytes // 1048576} MB > 50 MB)")
                    if len(members_data) >= _MAX_FILES:
                        raise ValueError(f"Số lượng file trong tệp nén vượt quá giới hạn ({_MAX_FILES})")

                    raw_data = full_path.read_bytes()
                    members_data.append((rel_name, f_size, raw_data))

            return members_data

    @staticmethod
    def _is_safe_path(name: str) -> bool:
        """
        Guards against Zip Slip (path traversal) vulnerabilities cross-platform.
        Ensures the relative archive path does not escape using '..' or absolute paths.
        """
        p = Path(name)
        return not p.is_absolute() and ".." not in p.parts and not name.startswith(("/", "\\"))

    @classmethod
    def _unpack_via_python_libs(
        cls,
        file_bytes: bytes,
        filename: str,
        mime_type: Optional[str] = None,
        password: Optional[str] = None,
    ) -> list[tuple[str, int, bytes]]:
        """
        Pure Python fallback extraction using pyzipper/zipfile, rarfile, py7zr, tarfile.
        """
        ext = Path(filename).suffix.lower()
        mime = (mime_type or "").lower()
        fmt = cls._detect_archive_type(ext, mime, file_bytes)

        _MAX_TOTAL_BYTES = 50 * 1024 * 1024
        _MAX_FILES = 200
        _MAX_SINGLE_FILE = 10 * 1024 * 1024

        members_data: list[tuple[str, int, bytes]] = []

        if fmt == "zip":
            zip_cls = None
            try:
                import pyzipper
                zip_cls = pyzipper.AESZipFile
            except ImportError:
                import zipfile
                zip_cls = zipfile.ZipFile

            try:
                with zip_cls(io.BytesIO(file_bytes)) as zf:
                    pwd_bytes = password.encode("utf-8") if password else None
                    if pwd_bytes:
                        zf.setpassword(pwd_bytes)

                    infolist = zf.infolist()
                    has_encrypted = any((m.flag_bits & 0x1) for m in infolist)
                    if has_encrypted and not password:
                        raise ArchivePasswordRequiredError(f"Tệp nén `{filename}` được đặt mật khẩu bảo vệ.")

                    total_uncompressed = sum(m.file_size for m in infolist)
                    if total_uncompressed > _MAX_TOTAL_BYTES:
                        raise ValueError(f"Dung lượng giải nén ZIP quá lớn ({total_uncompressed // 1048576} MB > 50 MB)")
                    if len(infolist) > _MAX_FILES:
                        raise ValueError(f"Số lượng file trong ZIP vượt quá giới hạn ({len(infolist)} > {_MAX_FILES})")

                    for info in infolist:
                        name = info.filename
                        if name.endswith("/"):
                            continue
                        if not cls._is_safe_path(name):
                            continue
                        if info.file_size > _MAX_SINGLE_FILE:
                            continue
                        try:
                            raw = zf.read(name, pwd=pwd_bytes)
                        except (RuntimeError, Exception) as re:
                            err_str = str(re).lower()
                            if "password required" in err_str or "encrypted" in err_str:
                                if not password:
                                    raise ArchivePasswordRequiredError(f"Tệp nén `{filename}` được đặt mật khẩu bảo vệ.")
                                raise ArchiveInvalidPasswordError(f"Mật khẩu '{password}' không chính xác cho tệp `{filename}`.")
                            if "bad password" in err_str or "crc" in err_str or "bad crc" in err_str:
                                raise ArchiveInvalidPasswordError(f"Mật khẩu '{password}' không chính xác cho tệp `{filename}`.")
                            raise

                        members_data.append((name, info.file_size, raw))
            except (ArchivePasswordRequiredError, ArchiveInvalidPasswordError):
                raise
            except Exception as ze:
                err_s = str(ze).lower()
                if "password" in err_s or "encrypted" in err_s:
                    if not password:
                        raise ArchivePasswordRequiredError(f"Tệp nén `{filename}` được đặt mật khẩu bảo vệ.")
                    raise ArchiveInvalidPasswordError(f"Mật khẩu '{password}' không chính xác cho tệp `{filename}`.")
                raise ArchiveCorruptedError(f"Không thể đọc tệp ZIP `{filename}`: {ze}")

        elif fmt == "rar":
            try:
                import rarfile
            except ImportError:
                raise ImportError("Thư viện 'rarfile' chưa được cài đặt trong hệ thống")
            try:
                with rarfile.RarFile(io.BytesIO(file_bytes)) as rf:
                    if rf.needs_password() and not password:
                        raise ArchivePasswordRequiredError(f"Tệp nén RAR `{filename}` được đặt mật khẩu bảo vệ.")

                    if password:
                        rf.setpassword(password)

                    infolist = rf.infolist()
                    total_uncompressed = sum(getattr(m, "file_size", 0) for m in infolist)
                    if total_uncompressed > _MAX_TOTAL_BYTES:
                        raise ValueError(f"Dung lượng giải nén RAR quá lớn ({total_uncompressed // 1048576} MB > 50 MB)")
                    if len(infolist) > _MAX_FILES:
                        raise ValueError(f"Số lượng file trong RAR vượt quá giới hạn ({len(infolist)} > {_MAX_FILES})")

                    for info in infolist:
                        name = info.filename
                        if getattr(info, "is_dir", lambda: False)():
                            continue
                        if not cls._is_safe_path(name):
                            continue
                        size = getattr(info, "file_size", 0)
                        if size > _MAX_SINGLE_FILE:
                            continue
                        try:
                            raw = rf.read(name, pwd=password)
                        except rarfile.PasswordRequired:
                            raise ArchivePasswordRequiredError(f"Tệp nén RAR `{filename}` được đặt mật khẩu bảo vệ.")
                        except rarfile.BadRarFile as bre:
                            if password:
                                raise ArchiveInvalidPasswordError(f"Mật khẩu '{password}' không chính xác cho tệp `{filename}`.")
                            raise ArchiveCorruptedError(f"Lỗi đọc RAR: {bre}")
                        members_data.append((name, size, raw))
            except (ArchivePasswordRequiredError, ArchiveInvalidPasswordError, ArchiveCorruptedError):
                raise
            except Exception as re:
                err_s = str(re).lower()
                if "password" in err_s:
                    if not password:
                        raise ArchivePasswordRequiredError(f"Tệp nén RAR `{filename}` được đặt mật khẩu bảo vệ.")
                    raise ArchiveInvalidPasswordError(f"Mật khẩu '{password}' không chính xác.")
                raise ArchiveCorruptedError(f"Lỗi đọc tệp RAR: {re}")

        elif fmt == "7z":
            try:
                import py7zr
            except ImportError:
                raise ImportError("Thư viện 'py7zr' chưa được cài đặt trong hệ thống")
            try:
                with py7zr.SevenZipFile(io.BytesIO(file_bytes), mode="r", password=password) as sz:
                    if sz.needs_password() and not password:
                        raise ArchivePasswordRequiredError(f"Tệp nén 7Z `{filename}` được đặt mật khẩu bảo vệ.")

                    archive_dict = sz.readall()
                    for name, bio in archive_dict.items():
                        if not cls._is_safe_path(name):
                            continue
                        raw = bio.getvalue() if hasattr(bio, "getvalue") else bio.read()
                        members_data.append((name, len(raw), raw))
            except (ArchivePasswordRequiredError, ArchiveInvalidPasswordError):
                raise
            except py7zr.exceptions.PasswordRequired:
                raise ArchivePasswordRequiredError(f"Tệp nén 7Z `{filename}` được đặt mật khẩu bảo vệ.")
            except py7zr.exceptions.Bad7zFile:
                if password:
                    raise ArchiveInvalidPasswordError(f"Mật khẩu '{password}' không chính xác cho tệp `{filename}`.")
                raise ArchiveCorruptedError(f"Tệp 7Z `{filename}` bị lỗi hoặc yêu cầu mật khẩu.")
            except Exception as se:
                err_s = str(se).lower()
                if "password" in err_s:
                    if not password:
                        raise ArchivePasswordRequiredError(f"Tệp nén 7Z `{filename}` được đặt mật khẩu bảo vệ.")
                    raise ArchiveInvalidPasswordError(f"Mật khẩu '{password}' không chính xác.")
                raise ArchiveCorruptedError(f"Lỗi đọc 7Z: {se}")

        elif fmt == "tar":
            import tarfile
            with tarfile.open(fileobj=io.BytesIO(file_bytes), mode="r:*") as tf:
                infolist = tf.getmembers()
                total_uncompressed = sum(m.size for m in infolist)
                if total_uncompressed > _MAX_TOTAL_BYTES:
                    raise ValueError(f"Dung lượng giải nén TAR quá lớn ({total_uncompressed // 1048576} MB > 50 MB)")
                if len(infolist) > _MAX_FILES:
                    raise ValueError(f"Số lượng file trong TAR vượt quá giới hạn ({len(infolist)} > {_MAX_FILES})")

                for info in infolist:
                    if info.isdir():
                        continue
                    name = info.name
                    if not cls._is_safe_path(name):
                        continue
                    if info.size > _MAX_SINGLE_FILE:
                        continue
                    extracted = tf.extractfile(info)
                    if extracted:
                        raw = extracted.read()
                        members_data.append((name, info.size, raw))

        else:
            raise ValueError(f"Định dạng tệp nén không được hỗ trợ: {ext} (MIME: {mime})")

        return members_data

    @classmethod
    def _unpack_archive_members(
        cls,
        file_bytes: bytes,
        filename: str,
        mime_type: Optional[str] = None,
        password: Optional[str] = None,
    ) -> list[tuple[str, int, bytes]]:
        """
        Safely unpacks all members of an archive in memory.
        Enforces security bounds:
        - Zip Slip guard (path traversal)
        - Zip Bomb guard (< 50MB uncompressed, max 200 files)
        - Password support for encrypted archives
        Returns: list of (member_name, file_size, raw_bytes)
        """
        if shutil.which("7z") or shutil.which("7za"):
            try:
                return cls._unpack_via_7z(file_bytes, filename, password=password)
            except (ArchivePasswordRequiredError, ArchiveInvalidPasswordError):
                raise
            except Exception as e7z:
                logger.warning("[MediaProcessor] 7z CLI extraction error (%s), falling back to Python libraries", e7z)

        return cls._unpack_via_python_libs(file_bytes, filename, mime_type, password=password)

    async def process_archive(
        self,
        file_bytes: bytes,
        mime_type: Optional[str],
        filename: str,
        caption: str = "",
        password: Optional[str] = None,
    ) -> str:
        """
        Universal Multi-modal Archive Extractor with Password Support.
        Recursively extracts:
        - Documents (PDF, DOCX, XLSX, CSV, JSON)
        - Source code and plain text files
        - Images via concise Technical Vision API (is_archive=True)
        - Reports full inventory manifest
        """
        # Auto-extract password from caption if not explicitly provided
        effective_pwd = password or extract_password_from_text(caption)

        # Let ArchivePasswordRequiredError and ArchiveInvalidPasswordError propagate to caller
        members = self._unpack_archive_members(file_bytes, filename, mime_type, password=effective_pwd)

        if not members:
            return f"(Tệp nén `{filename}` trống hoặc không chứa file hợp lệ.)"

        _MAX_IMAGES = 3
        images_analyzed = 0

        manifest_lines: List[str] = []
        detail_blocks: List[str] = []

        for idx, (m_name, m_size, m_data) in enumerate(members, 1):
            icon, type_label, text_content = self._extract_child_file_content(m_name, m_data)
            size_kb = round(m_size / 1024, 1)

            # Record in manifest
            manifest_lines.append(f"  {idx}. {icon} `{m_name}` ({size_kb} KB — {type_label})")

            # Extract detailed content
            if text_content is not None:
                detail_blocks.append(f"━━━ [{idx}/{len(members)}] {icon} `{m_name}` ({type_label}) ━━━\n{text_content}")
            elif icon == "🖼️":
                if images_analyzed < _MAX_IMAGES:
                    try:
                        logger.info("[MediaProcessor] Vision analysis for archive image: %s (%d bytes)", m_name, m_size)
                        v_result = await self.analyze_image(
                            m_data,
                            caption=caption or f"Tệp ảnh {m_name} trong archive {filename}",
                            is_archive=True,
                        )
                        detail_blocks.append(f"━━━ [{idx}/{len(members)}] 🖼️ `{m_name}` ({type_label}) ━━━\n{v_result}")
                        images_analyzed += 1
                    except Exception as exc:
                        detail_blocks.append(f"━━━ [{idx}/{len(members)}] 🖼️ `{m_name}` ━━━\n(Lỗi phân tích hình ảnh: {exc})")
                else:
                    detail_blocks.append(f"━━━ [{idx}/{len(members)}] 🖼️ `{m_name}` ━━━\n(Bỏ qua phân tích thị giác do đã đạt giới hạn {_MAX_IMAGES} ảnh/tệp nén)")
            else:
                detail_blocks.append(f"━━━ [{idx}/{len(members)}] ⚙️ `{m_name}` ({type_label}) ━━━\n[Tệp nhị phân không khả dụng trích xuất văn bản/hình ảnh]")

        header = (
            f"[📦 TỔNG QUAN TỆP NÉN: '{filename}']\n"
            f"• Tổng số tệp trích xuất: {len(members)} tệp ({images_analyzed} hình ảnh đã phân tích thị giác)\n"
            f"• Danh mục tệp bên trong:\n" + "\n".join(manifest_lines) + "\n\n"
            f"═══════════════════════════════════════════════════════════\n"
            f"[CHI TIẾT NỘI DUNG ĐÃ TRÍCH XUẤT TỪNG TỆP BÊN DƯỚI]\n\n"
        )

        full_result = header + "\n\n".join(detail_blocks)

        if len(full_result) > _MAX_DOC_CHARS:
            full_result = full_result[:_MAX_DOC_CHARS] + f"\n\n... (nội dung cắt bớt, chỉ hiển thị {_MAX_DOC_CHARS} ký tự đầu)"

        return full_result

    def extract_document_text(
        self,
        file_bytes: bytes,
        mime_type: Optional[str],
        filename: str,
    ) -> str:
        """
        Standalone document extractor for single files (PDF, Word, Excel, CSV, JSON, Text/Code).
        """
        icon, type_label, text_content = self._extract_child_file_content(filename, file_bytes)
        if text_content is not None:
            if len(text_content) > _MAX_DOC_CHARS:
                text_content = text_content[:_MAX_DOC_CHARS] + f"\n\n... (nội dung cắt bớt, chỉ hiển thị {_MAX_DOC_CHARS} ký tự đầu)"
            return text_content

        ext = Path(filename).suffix.lower()
        return (
            f"(Định dạng tệp `{ext or 'không rõ'}` ({type_label}) chưa được hỗ trợ đọc trực tiếp. "
            "Hỗ trợ: PDF, DOCX, XLSX, CSV, JSON, ZIP, RAR, TAR, 7Z và các tệp mã nguồn.)"
        )
